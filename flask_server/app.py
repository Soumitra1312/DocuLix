from flask import Flask, request, jsonify, url_for
from paraphrase import paraphrase
from predict import run_prediction
from io import StringIO
import json
from flask_cors import CORS
import PyPDF2
import requests
from dotenv import load_dotenv
import os
import re
import numpy as np
import hashlib
import time
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor
import threading

# Import libraries for document and image processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available, Word document support disabled")

try:
    from PIL import Image
    import pytesseract
    import easyocr
    OCR_AVAILABLE = True
    # Initialize EasyOCR reader
    ocr_reader = easyocr.Reader(['en'])
except ImportError:
    OCR_AVAILABLE = False
    print("OCR libraries not available, image text extraction disabled")

# GPU acceleration disabled - CuPy removed from dependencies
# Keep GPU_AVAILABLE flag for compatibility with existing code paths
GPU_AVAILABLE = False

load_dotenv()

app = Flask(__name__)
# Allow CORS for both localhost and 127.0.0.1 on port 3000 (Next.js dev)
CORS(app, origins=["http://localhost:3000", "http://127.0.0.1:3000"], supports_credentials=True)
answers = []

# Cache for storing processed PDF chunks
pdf_cache = {}
cache_lock = threading.Lock()

class PDFCache:
    def __init__(self):
        self.cache = {}
        self.cache_timestamps = {}
        self.max_cache_age = timedelta(hours=24)  # Cache for 24 hours
        self.lock = threading.Lock()
    
    def get_file_hash(self, file_content):
        """Generate hash for file content"""
        return hashlib.md5(file_content).hexdigest()
    
    def is_cached(self, file_hash):
        """Check if file chunks are cached and not expired"""
        with self.lock:
            if file_hash not in self.cache:
                return False
            
            # Check if cache is expired
            if datetime.now() - self.cache_timestamps[file_hash] > self.max_cache_age:
                del self.cache[file_hash]
                del self.cache_timestamps[file_hash]
                return False
            
            return True
    
    def get_chunks(self, file_hash):
        """Get cached chunks"""
        with self.lock:
            return self.cache.get(file_hash, {})
    
    def store_chunks(self, file_hash, text, chunks, simplified_chunks=None):
        """Store processed chunks in cache"""
        with self.lock:
            self.cache[file_hash] = {
                'text': text,
                'chunks': chunks,
                'simplified_chunks': simplified_chunks or chunks,
                'processed_at': datetime.now()
            }
            self.cache_timestamps[file_hash] = datetime.now()
    
    def clear_expired(self):
        """Clear expired cache entries"""
        with self.lock:
            current_time = datetime.now()
            expired_keys = [
                key for key, timestamp in self.cache_timestamps.items()
                if current_time - timestamp > self.max_cache_age
            ]
            for key in expired_keys:
                del self.cache[key]
                del self.cache_timestamps[key]

# Initialize PDF cache
pdf_cache = PDFCache()

# Configure OpenRouter API
OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY')
OPENROUTER_BASE_URL = os.getenv('OPENROUTER_BASE_URL', 'https://openrouter.ai/api/v1')
CURRENT_MODEL = os.getenv('CURRENT_MODEL', 'google/gemini-flash-1.5')

# Available models for fallback
AVAILABLE_MODELS = [
    'google/gemini-flash-1.5',
    'google/gemini-pro-1.5',
    'openai/gpt-4o-mini',
    'openai/gpt-3.5-turbo'
]

if OPENROUTER_API_KEY:
    print(f"OpenRouter API configured with model: {CURRENT_MODEL}")
else:
    print("Warning: OPENROUTER_API_KEY not found in environment variables")

def call_openrouter_api(prompt, model=None, max_tokens=4000, temperature=0.7):
    """
    Call OpenRouter API with the given prompt
    """
    if not OPENROUTER_API_KEY:
        raise ValueError("OpenRouter API key not configured")
    
    # Use the specified model or fall back to the current model
    model_to_use = model or CURRENT_MODEL
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:5000",  # Optional, for rankings
        "X-Title": "Legal AI Assistant"  # Optional, for rankings
    }
    
    data = {
        "model": model_to_use,
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": temperature
    }
    
    try:
        response = requests.post(
            f"{OPENROUTER_BASE_URL}/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            if 'choices' in result and len(result['choices']) > 0:
                return result['choices'][0]['message']['content']
            else:
                raise ValueError("No response content received from OpenRouter API")
        else:
            # Try with a fallback model if the current model fails
            if model_to_use in AVAILABLE_MODELS and len(AVAILABLE_MODELS) > 1:
                fallback_models = [m for m in AVAILABLE_MODELS if m != model_to_use]
                for fallback_model in fallback_models:
                    try:
                        data["model"] = fallback_model
                        fallback_response = requests.post(
                            f"{OPENROUTER_BASE_URL}/chat/completions",
                            headers=headers,
                            json=data,
                            timeout=60
                        )
                        if fallback_response.status_code == 200:
                            result = fallback_response.json()
                            if 'choices' in result and len(result['choices']) > 0:
                                print(f"Used fallback model: {fallback_model}")
                                return result['choices'][0]['message']['content']
                    except Exception as e:
                        continue
            
            raise ValueError(f"OpenRouter API error: {response.status_code} - {response.text}")
            
    except requests.exceptions.Timeout:
        raise ValueError("Request to OpenRouter API timed out")
    except requests.exceptions.RequestException as e:
        raise ValueError(f"Request to OpenRouter API failed: {str(e)}")

def list_available_models():
    """
    Get list of available models from OpenRouter
    """
    if not OPENROUTER_API_KEY:
        return AVAILABLE_MODELS
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    
    try:
        response = requests.get(
            f"{OPENROUTER_BASE_URL}/models",
            headers=headers,
            timeout=10
        )
        
        if response.status_code == 200:
            models_data = response.json()
            return [model['id'] for model in models_data.get('data', [])]
        else:
            return AVAILABLE_MODELS
    except:
        return AVAILABLE_MODELS



def validate_legal_document_with_ai(text_chunks):
    """
    Use AI to determine if the document is a legal document
    Returns (is_legal, confidence, document_type, explanation)
    """
    if not OPENROUTER_API_KEY:
        print("⚠️ OpenRouter API key not available, skipping AI validation")
        return False, 0.2, "Unknown", "API validation unavailable - cannot verify document type"
    
    # Use first 3 chunks for analysis (or all if less than 3)
    sample_text = " ".join(text_chunks[:3]) if len(text_chunks) >= 3 else " ".join(text_chunks)
    
    # Limit text length for API efficiency
    if len(sample_text) > 4000:
        sample_text = sample_text[:4000] + "..."
    
    validation_prompt = f"""
You are a legal document classifier. Analyze the following document content and determine if this is a legal document.

Document Content:
{sample_text}

STRICT CRITERIA: Only classify as legal if the document contains:
- Legal terminology, clauses, or contractual language
- Official legal formatting or structure
- Legal obligations, rights, or duties
- Regulatory or statutory content
- Court-related or judicial content

NON-LEGAL documents include:
- General business documents, reports, manuals
- Academic papers, research documents
- Personal letters, emails, or correspondence  
- Technical documentation, user guides
- Marketing materials, brochures
- News articles, blog posts
- Financial statements (unless legally binding)

Respond in this exact JSON format:
{{
    "is_legal_document": true/false,
    "confidence": 0.0-1.0,
    "document_type": "specific type or 'Non-legal'",
    "explanation": "brief explanation of your decision"
}}

Be STRICT - when in doubt, classify as non-legal. Respond ONLY with the JSON, no other text.
"""

    try:
        response = call_openrouter_api(validation_prompt, max_tokens=300, temperature=0.0)
        
        # Parse JSON response
        import json
        try:
            # Clean the response to extract JSON
            response_clean = response.strip()
            if response_clean.startswith('```json'):
                response_clean = response_clean.replace('```json', '').replace('```', '').strip()
            
            result = json.loads(response_clean)
            is_legal = result.get('is_legal_document', False)
            confidence = float(result.get('confidence', 0.0))
            
            # Fix confidence if it's already in percentage format (>1.0)
            if confidence > 1.0:
                confidence = confidence / 100.0
                
            # Ensure confidence is between 0.0 and 1.0
            confidence = max(0.0, min(1.0, confidence))
            
            doc_type = result.get('document_type', 'Unknown')
            explanation = result.get('explanation', 'No explanation provided')
            
            return is_legal, confidence, doc_type, explanation
            
        except json.JSONDecodeError:
            print(f"⚠️ Failed to parse AI validation response: {response}")
            # More conservative fallback - assume non-legal unless clear indicators
            response_lower = response.lower()
            if any(term in response_lower for term in ['contract', 'agreement', 'legal', 'terms', 'clause']):
                return True, 0.6, "Legal Document", "Fallback validation - found legal indicators"
            else:
                return False, 0.8, "Non-legal Document", "Fallback validation - no clear legal indicators"
                
    except Exception as e:
        print(f"⚠️ AI validation error: {e}")
        # Conservative fallback - assume non-legal if we can't validate
        return False, 0.5, "Unknown", f"Validation failed: {str(e)}"

def load_questions_short():
    # Define 10 most common predefined questions
    predefined_questions = [
        "What is the contract name?",
        "Who are the parties that signed the contract?",
        "What is the agreement date of the contract?",
        "What is the date when the contract is effective?",
        "What date will the contract's initial term expire?",
        "What is the renewal term after the initial term expires?",
        "What is the notice period required to terminate renewal?",
        "Which state/country's law governs the interpretation of the contract?",
        "Can a party terminate this contract without cause?",
        "What are the payment terms and conditions?"
    ]
    
    return predefined_questions


# OpenRouter API function removed - using direct Gemini API instead


def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        print(f"DEBUG: PDF extraction - file position before: {file.tell()}")
        pdf_reader = PyPDF2.PdfReader(file)
        print(f"DEBUG: PDF reader created, pages: {len(pdf_reader.pages)}")
        text = ""
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            print(f"DEBUG: Page {i+1} text length: {len(page_text)}")
            text += page_text + "\n"
        print(f"DEBUG: Total PDF text length: {len(text)}")
        return text
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return None

def extract_text_from_docx(file):
    """Extract text from Word document (DOCX)"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text if text.strip() else None
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return None

def extract_text_from_image(file):
    """Extract text from image using OCR"""
    if not OCR_AVAILABLE:
        return None
    
    try:
        # Open image with PIL
        image = Image.open(file)
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Try EasyOCR first (generally more accurate)
        try:
            # Reset file pointer for EasyOCR
            file.seek(0)
            results = ocr_reader.readtext(file.read())
            text = " ".join([result[1] for result in results])
            if text.strip():
                return text
        except Exception as e:
            print(f"EasyOCR failed: {e}")
        
        # Fallback to Tesseract
        try:
            # Reset file pointer and reopen image for Tesseract
            file.seek(0)
            image = Image.open(file)
            text = pytesseract.image_to_string(image)
            if text.strip():
                return text
        except Exception as e:
            print(f"Tesseract OCR failed: {e}")
        
        return None
    except Exception as e:
        print(f"Error extracting text from image: {e}")
        return None

def get_file_type(filename):
    """Determine file type from filename"""
    if not filename:
        return 'unknown'
    
    extension = filename.lower().split('.')[-1]
    
    if extension == 'pdf':
        return 'pdf'
    elif extension in ['doc', 'docx']:
        return 'word'
    elif extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp']:
        return 'image'
    elif extension in ['txt']:
        return 'text'
    else:
        return 'unknown'

def extract_text_from_file(file, file_type):
    """Extract text from file based on its type"""
    if file_type == 'pdf':
        return extract_text_from_pdf(file)
    elif file_type == 'word':
        return extract_text_from_docx(file)
    elif file_type == 'image':
        return extract_text_from_image(file)
    elif file_type == 'text':
        # Handle text files with different encodings
        print(f"DEBUG: Text file extraction - file position before: {file.tell()}")
        file_content = file.read()
        print(f"DEBUG: Text file content length: {len(file_content)}")
        try:
            text_content = file_content.decode("utf-8")
            print(f"DEBUG: UTF-8 decode successful, length: {len(text_content)}")
            return text_content
        except UnicodeDecodeError as e:
            print(f"DEBUG: UTF-8 decode failed: {e}")
            try:
                text_content = file_content.decode("latin-1")
                print(f"DEBUG: Latin-1 decode successful, length: {len(text_content)}")
                return text_content
            except UnicodeDecodeError as e2:
                print(f"DEBUG: Latin-1 decode failed: {e2}")
                text_content = file_content.decode("cp1252")
                print(f"DEBUG: CP1252 decode successful, length: {len(text_content)}")
                return text_content
    else:
        return None

def is_legal_document(text_content):
    """
    Validate if the document content appears to be a legal document
    Returns (is_legal, confidence_score, detected_indicators)
    """
    if not text_content or len(text_content.strip()) < 100:
        return False, 0.0, []
    
    # Convert to lowercase for analysis
    text_lower = text_content.lower()
    
    # Legal keywords and phrases categorized by strength
    strong_legal_indicators = [
        'whereas', 'wherefore', 'heretofore', 'hereinafter', 'party of the first part',
        'party of the second part', 'in witness whereof', 'terms and conditions',
        'force and effect', 'null and void', 'ipso facto', 'prima facie',
        'res ipsa loquitur', 'quid pro quo', 'habeas corpus', 'amicus curiae',
        'statutory', 'jurisdiction', 'litigation', 'plaintiff', 'defendant',
        'appellant', 'appellee', 'breach of contract', 'damages', 'injunction',
        'cease and desist', 'intellectual property', 'copyright', 'trademark',
        'patent', 'confidentiality agreement', 'non-disclosure', 'indemnification',
        'liability', 'negligence', 'tort', 'covenant', 'warranty', 'representation',
        'arbitration', 'mediation', 'settlement', 'judgment', 'decree',
        'subpoena', 'deposition', 'affidavit', 'exhibit', 'evidence'
    ]
    
    medium_legal_indicators = [
        'contract', 'agreement', 'legal', 'law', 'clause', 'section', 'provision',
        'article', 'amendment', 'regulation', 'statute', 'code', 'act',
        'compliance', 'violation', 'breach', 'obligation', 'right', 'duty',
        'consent', 'authorization', 'license', 'permit', 'certificate',
        'court', 'judge', 'jury', 'trial', 'hearing', 'proceeding',
        'motion', 'brief', 'pleading', 'complaint', 'answer', 'counter-claim',
        'discovery', 'interrogatory', 'admission', 'penalty', 'fine',
        'sanction', 'punishment', 'sentence', 'probation', 'parole'
    ]
    
    weak_legal_indicators = [
        'shall', 'may', 'must', 'required', 'prohibited', 'permitted',
        'entitled', 'responsible', 'accountable', 'binding', 'enforceable',
        'effective', 'terminate', 'expire', 'renew', 'modify', 'amend',
        'notify', 'inform', 'disclose', 'confidential', 'proprietary',
        'ownership', 'title', 'interest', 'benefit', 'consideration',
        'payment', 'compensation', 'fee', 'cost', 'expense'
    ]
    
    # Document structure indicators
    structure_indicators = [
        'article i', 'article 1', 'section 1', 'section i', 'clause',
        'subsection', 'paragraph', 'subparagraph', 'exhibit a', 'exhibit 1',
        'schedule a', 'schedule 1', 'appendix a', 'appendix 1',
        'witnesseth', 'recitals', 'definitions', 'interpretation'
    ]
    
    # Legal document types
    document_types = [
        'lease agreement', 'rental agreement', 'employment contract',
        'service agreement', 'purchase agreement', 'sales contract',
        'partnership agreement', 'shareholders agreement', 'merger agreement',
        'acquisition agreement', 'licensing agreement', 'franchise agreement',
        'joint venture agreement', 'non-compete agreement', 'severance agreement',
        'settlement agreement', 'plea agreement', 'divorce decree',
        'custody agreement', 'will and testament', 'trust agreement',
        'power of attorney', 'mortgage', 'deed', 'title', 'lien',
        'security agreement', 'promissory note', 'loan agreement',
        'credit agreement', 'insurance policy', 'warranty',
        'terms of service', 'privacy policy', 'user agreement',
        'software license', 'copyright license', 'trademark license',
        'patent license', 'assignment agreement', 'transfer agreement'
    ]
    
    # Count indicators
    strong_count = sum(1 for indicator in strong_legal_indicators if indicator in text_lower)
    medium_count = sum(1 for indicator in medium_legal_indicators if indicator in text_lower)
    weak_count = sum(1 for indicator in weak_legal_indicators if indicator in text_lower)
    structure_count = sum(1 for indicator in structure_indicators if indicator in text_lower)
    doc_type_count = sum(1 for doc_type in document_types if doc_type in text_lower)
    
    # Calculate weighted score
    score = (
        strong_count * 10 +
        medium_count * 5 +
        weak_count * 2 +
        structure_count * 8 +
        doc_type_count * 15
    )
    
    # Normalize score based on document length
    words_count = len(text_content.split())
    normalized_score = min(score / max(words_count / 100, 1), 100)
    
    # Determine if it's a legal document - lowered threshold for more lenient validation
    is_legal = normalized_score >= 8  # More lenient threshold for legal document
    
    # Collect detected indicators for feedback
    detected_indicators = []
    if strong_count > 0:
        detected_indicators.extend([ind for ind in strong_legal_indicators if ind in text_lower][:3])
    if medium_count > 0:
        detected_indicators.extend([ind for ind in medium_legal_indicators if ind in text_lower][:3])
    if doc_type_count > 0:
        detected_indicators.extend([doc_type for doc_type in document_types if doc_type in text_lower][:2])
    
    confidence = min(normalized_score / 20, 1.0)  # Convert to 0-1 scale with lower denominator
    
    return is_legal, confidence, detected_indicators[:5]  # Return top 5 indicators

def chunk_text_gpu_optimized(text, chunk_size=2000, overlap=200):
    """Optimized GPU-accelerated text chunking using CuPy with parallel processing"""
    if not GPU_AVAILABLE:
        return chunk_text_cpu_parallel(text, chunk_size, overlap)
    
    try:
        if len(text) <= chunk_size:
            return [text]
        
        start_time = time.time()
        print(f"Starting GPU chunking for text of length {len(text)}")
        
        # Pre-compute chunk boundaries for parallel processing
        text_len = len(text)
        chunk_starts = []
        start = 0
        
        while start < text_len:
            chunk_starts.append(start)
            start += chunk_size - overlap
        
        # Process chunks in parallel using GPU
        chunks = []
        batch_size = min(10, len(chunk_starts))  # Process in batches to manage GPU memory
        
        for i in range(0, len(chunk_starts), batch_size):
            batch_starts = chunk_starts[i:i + batch_size]
            batch_chunks = []
            
            for start in batch_starts:
                end = min(start + chunk_size, text_len)
                chunk = text[start:end]
                
                # Try to end chunk at a sentence boundary
                if end < text_len:
                    sentence_end = chunk.rfind('.')
                    if sentence_end > chunk_size * 0.7:
                        chunk = chunk[:sentence_end + 1]
                
                batch_chunks.append(chunk.strip())
            
            chunks.extend(batch_chunks)
            
            # Remove arbitrary limit - process entire document
        
        end_time = time.time()
        print(f"GPU chunking completed in {end_time - start_time:.2f}s, created {len(chunks)} chunks")
        
        return chunks
        
    except Exception as e:
        print(f"GPU chunking failed: {e}, falling back to CPU")
        return chunk_text_cpu_parallel(text, chunk_size, overlap)

def chunk_text_cpu_parallel(text, chunk_size=2000, overlap=200):
    """CPU-based parallel text chunking"""
    if len(text) <= chunk_size:
        return [text]
    
    start_time = time.time()
    print(f"Starting CPU parallel chunking for text of length {len(text)}")
    
    # Pre-compute chunk boundaries
    text_len = len(text)
    chunk_starts = []
    start = 0
    
    while start < text_len:
        chunk_starts.append(start)
        start += chunk_size - overlap
    
    def process_chunk(start):
        end = min(start + chunk_size, text_len)
        chunk = text[start:end]
        
        # Try to end chunk at a sentence boundary
        if end < text_len:
            sentence_end = chunk.rfind('.')
            if sentence_end > chunk_size * 0.7:
                chunk = chunk[:sentence_end + 1]
        
        return chunk.strip()
    
    # Use ThreadPoolExecutor for parallel processing - process ALL chunks
    with ThreadPoolExecutor(max_workers=4) as executor:
        chunks = list(executor.map(process_chunk, chunk_starts))  # Process entire document
    
    end_time = time.time()
    print(f"CPU parallel chunking completed in {end_time - start_time:.2f}s, created {len(chunks)} chunks")
    
    return chunks

def chunk_text_gpu(text, chunk_size=2000, overlap=200):
    """Legacy GPU chunking function - kept for compatibility"""
    return chunk_text_gpu_optimized(text, chunk_size, overlap)

def chunk_text_cpu(text, chunk_size=2000, overlap=200):
    """CPU-based text chunking (fallback)"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to end chunk at a sentence boundary
        if end < len(text):
            sentence_end = chunk.rfind('.')
            if sentence_end > chunk_size * 0.7:  # Only if we find a sentence end in the latter part
                chunk = chunk[:sentence_end + 1]
                end = start + sentence_end + 1
        
        chunks.append(chunk.strip())
        start = end - overlap
        
        if start >= len(text):
            break
    
    return chunks

def chunk_text(text, chunk_size=3000, overlap=300):
    """Fast chunking function with larger chunks for better performance"""
    if len(text) <= chunk_size:
        return [text]
    
    print("Using fast chunking mode")
    
    # Simple fast chunking without GPU overhead for better speed
    chunks = []
    start = 0
    chunk_count = 0
    
    while start < len(text) and chunk_count < 20:  # Limit chunks for speed
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to end chunk at a sentence boundary
        if end < len(text):
            sentence_end = chunk.rfind('.')
            if sentence_end > chunk_size * 0.7:
                chunk = chunk[:sentence_end + 1]
                end = start + sentence_end + 1
        
        chunks.append(chunk.strip())
        start = end - overlap
        chunk_count += 1
        
        if start >= len(text):
            break
    
    print(f"Fast chunking completed: created {len(chunks)} chunks")
    return chunks

# Available models function removed - using Gemini API directly

def find_relevant_chunks(question, text_chunks):
    """Find chunks most relevant to the question using advanced keyword matching and scoring"""
    import re
    
    # Extract key terms from the question
    question_lower = question.lower()
    
    # Score ALL chunks for complete document coverage
    chunk_scores = []
    
    print(f"Analyzing {len(text_chunks)} chunks for relevance...")
    
    for i, chunk in enumerate(text_chunks):
        chunk_lower = chunk.lower()
        score = 0
        
        # Extract key terms from question (remove common words)
        question_words = re.findall(r'\b\w+\b', question_lower)
        key_terms = [word for word in question_words if len(word) > 2 and 
                    word not in ['what', 'when', 'where', 'which', 'that', 'this', 'they', 'with', 'from', 'have', 'been', 'will', 'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'its', 'may', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 'did', 'she', 'use', 'her', 'now', 'him', 'way']]
        
        # Enhanced scoring system
        for term in key_terms:
            if term in chunk_lower:
                # Weight by frequency and term importance
                frequency = chunk_lower.count(term)
                if len(term) > 4:  # Longer terms are more specific
                    score += frequency * 3
                else:
                    score += frequency * 2
        
        # Look for legal section references (high priority)
        section_matches = re.findall(r'§\s*\d+(?:\([a-z]\))?', chunk)
        if section_matches:
            score += len(section_matches) * 10
            
        # Look for general section references
        if re.search(r'section\s+\d+', chunk_lower):
            score += 5
            
        # Look for subsection references  
        if re.search(r'\(\w+\)', chunk):
            score += 2
            
        # Enhanced legal term matching
        legal_terms = {
            'payment': 8, 'trustee': 10, 'compensation': 8, 'fee': 6, 
            'multiple': 5, 'assigned': 7, 'receives': 8, 'entitled': 7,
            'distribution': 6, 'allocation': 6, 'divided': 5, 'shared': 5,
            'court': 4, 'order': 4, 'approval': 5, 'bankruptcy': 8,
            'estate': 6, 'debtor': 6, 'creditor': 5, 'proceeding': 4
        }
        
        for term, weight in legal_terms.items():
            if term in question_lower and term in chunk_lower:
                score += chunk_lower.count(term) * weight
        
        # Contextual bonuses
        if 'multiple' in question_lower and 'trustee' in question_lower:
            if 'multiple' in chunk_lower and 'trustee' in chunk_lower:
                score += 15  # High relevance for multiple trustee questions
        
        # Boost score for chunks with numbers/amounts (often relevant to payment questions)
        if re.search(r'\$[\d,]+', chunk) or re.search(r'\b\d+\.\d+\b', chunk):
            score += 3
            
        # Boost for procedural language
        procedural_terms = ['shall', 'must', 'required', 'entitled', 'pursuant', 'accordance']
        for term in procedural_terms:
            if term in chunk_lower:
                score += 2
        
        chunk_scores.append((score, i, chunk))
    
    # Sort by score (highest first)
    chunk_scores.sort(key=lambda x: x[0], reverse=True)
    
    # Return ALL chunks in order of relevance (no filtering by score)
    # This ensures complete document coverage
    relevant_chunks = [chunk for score, i, chunk in chunk_scores]
    
    # Log scoring results
    high_score_count = len([s for s, i, c in chunk_scores if s > 0])
    print(f"Ranked all {len(relevant_chunks)} chunks by relevance ({high_score_count} with positive scores)")
    
    return relevant_chunks

def simplify_chunk_with_gemini(chunk):
    """Simplify a text chunk using OpenRouter API"""
    if not OPENROUTER_API_KEY:
        return chunk  # Return original chunk if API not available
    
    try:
        prompt = f"""
        Please simplify and extract the key legal information from the following text chunk. 
        Focus on important contractual details like parties, dates, terms, obligations, and conditions.
        Keep the essential information while removing unnecessary verbosity.
        
        Text to simplify:
        {chunk}
        
        Simplified version:
        """
        
        response = call_openrouter_api(prompt, max_tokens=2000, temperature=0.3)
        return response.strip()
    except Exception as e:
        print(f"Error simplifying chunk with OpenRouter: {e}")
        return chunk  # Return original chunk if simplification fails

def clean_ai_response(response):
    """Clean up AI response by removing only unnecessary meta-commentary while preserving detailed explanations"""
    if not response:
        return response
    
    # List of phrases to remove (case-insensitive) - remove disclaimers and meta-commentary
    phrases_to_remove = [
        "This explanation provides a general overview. Always consult the actual legal document for precise details and requirements.",
        "Always consult the actual legal document for precise details and requirements.",
        "Please consult the actual legal document for precise details.",
        "For precise details and requirements, always refer to the actual legal document.",
        "This is a general overview. Always consult the actual document for specific requirements.",
        "Please consult with a qualified legal professional",
        "Please consult a qualified legal professional",
        "Consult with a qualified legal professional",
        "Consult a qualified legal professional",
        "This is a general overview.",
        "This provides a general overview.",
        "For specific legal advice, please consult",
        "I am not a lawyer",
        "This is not legal advice",
        "The document provides a complete description of the process for calculating per-case compensation",
        "No additional information is needed from external sources to answer the question based solely on the provided text",
        "This information is based solely on the provided document sections",
        "A complete understanding would require access to the referenced documents"
    ]
    
    cleaned_response = response
    for phrase in phrases_to_remove:
        # Remove the phrase and any trailing periods/spaces
        import re
        pattern = re.compile(re.escape(phrase) + r'\.?\s*', re.IGNORECASE)
        cleaned_response = pattern.sub('', cleaned_response)
    
    # Remove markdown formatting that might appear
    markdown_patterns = [
        (r'\*\*(.*?)\*\*', lambda m: m.group(1).upper()),  # **bold** -> UPPERCASE
        (r'\*(.*?)\*', lambda m: m.group(1)),              # *italic* -> plain text
        (r'__(.*?)__', lambda m: m.group(1).upper()),      # __underline__ -> UPPERCASE  
        (r'_(.*?)_', lambda m: m.group(1)),                # _italic_ -> plain text
        (r'`(.*?)`', lambda m: m.group(1)),                # `code` -> plain text
        (r'\[(.*?)\]\(.*?\)', lambda m: m.group(1)),       # [text](link) -> text only
        (r'#{1,6}\s*', lambda m: ''),                      # ### headings -> remove hashes
    ]
    
    # Apply markdown removal
    for pattern, replacement in markdown_patterns:
        if callable(replacement):
            cleaned_response = re.sub(pattern, replacement, cleaned_response)
        else:
            cleaned_response = re.sub(pattern, replacement, cleaned_response)
    
    # Clean up extra whitespace and formatting while preserving structure
    cleaned_response = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned_response)  # Remove excessive newlines
    cleaned_response = cleaned_response.strip()
    
    return cleaned_response

def query_gemini(question, text_chunks, chunks_already_simplified=False):
    """Fast OpenRouter API query with optimized processing"""
    if not OPENROUTER_API_KEY:
        return "OpenRouter API not properly configured"
    
    try:
        print(f"Fast processing mode: analyzing {len(text_chunks)} chunks...")
        
        # Fast mode: Use top relevant chunks only for speed
        relevant_chunks = find_relevant_chunks(question, text_chunks)
        
        # Process top 12 most relevant chunks for detailed coverage
        top_chunks = relevant_chunks[:12]
        print(f"Using top {len(top_chunks)} most relevant chunks for detailed response")
        
        # Fast single query approach
        combined_text = "\n\n".join(top_chunks)
        
        print("Sending optimized query to OpenRouter...")
        
        prompt = f"""You are a legal expert explaining complex legal documents to everyday people. Based on the following legal document content, provide a detailed, easy-to-understand answer.

IMPORTANT FORMATTING RULES:
- Use CAPITAL LETTERS for headings and emphasis (not ** or markdown)
- Use dashes (-) or numbers (1., 2., 3.) for lists
- Use line breaks to separate sections clearly
- Use simple punctuation and spacing for emphasis
- Do NOT use **, [], or other markdown formatting

CONTENT INSTRUCTIONS:
1. Explain in Plain English: Use simple, everyday language that anyone can understand
2. Be Comprehensive: Provide detailed explanations with multiple aspects of the answer
3. Use Examples: When possible, include practical examples or scenarios
4. Break Down Complex Terms: Explain any legal jargon or technical terms in simple words
5. Structure Your Answer: Use clear paragraphs and numbered/bulleted lists for easy reading
6. Include Context: Explain why this information matters and how it applies in real situations
7. Reference Sources: Mention specific sections, subsections, or legal citations when available
8. Cover All Aspects: Address different scenarios, exceptions, or variations that might apply

Think of yourself as explaining this to a friend or family member who has no legal background.

Document Content:
{combined_text}

Question: {question}

Provide a detailed, comprehensive answer in plain text format (no markdown):"""
        
        response = call_openrouter_api(prompt, max_tokens=4000, temperature=0.7)
        answer = response.strip()
        
        if answer and not any(phrase in answer.lower() for phrase in [
            "not available", "not specified", "does not specify", "not mentioned", 
            "not found", "no information", "cannot determine"
        ]):
            print("Found answer in fast mode")
            return clean_ai_response(answer)
        else:
            return "The specific information requested was not found in the most relevant sections of the document."
    
    except Exception as e:
        print(f"Error querying OpenRouter: {e}")
        return f"Error processing question: {str(e)}. Please try again or check your API configuration."

# Function removed - no longer analyzing sentiment



questions_short = load_questions_short()



@app.route('/')
def index():
    """API health check endpoint"""
    return jsonify({"status": "ok", "message": "Legal AI Flask API is running", "version": "2.0"})

@app.route('/questionsshort')
def getQuestionsShort():
    return jsonify(questions_short)

@app.route('/models')
def get_available_models():
    """Debug route to check available OpenRouter models"""
    if not OPENROUTER_API_KEY:
        return json.dumps({"error": "API key not configured"})
    
    try:
        available_models = list_available_models()
        return json.dumps({"available_models": available_models, "current_model": CURRENT_MODEL})
    except Exception as e:
        return json.dumps({"error": str(e)})

@app.route('/test-api')
def test_openrouter_api():
    """Test OpenRouter API connection"""
    if not OPENROUTER_API_KEY:
        return json.dumps({"error": "API key not configured"})
    
    test_prompt = "Say 'Hello, this is a test from OpenRouter API!'"
    print(f" Testing OpenRouter API with simple prompt...")
    
    try:
        response = call_openrouter_api(test_prompt, max_tokens=100, temperature=0.1)
        return json.dumps({"success": True, "response": response})
    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})

@app.route('/upload-pdf', methods=["POST"])
def upload_and_process_pdf():
    """Upload and preprocess documents (PDF, Word, Images), return file hash for subsequent questions"""
    print("Document upload endpoint called")
    
    try:
        # Check if file is present in request
        if 'file' not in request.files:
            print("DEBUG: No 'file' key in request.files")
            return json.dumps({"error": "No file uploaded", "success": False})
        
        file = request.files["file"]
        print(f"DEBUG: File object: {file}")
        print(f"DEBUG: File filename: {file.filename}")
        
        if file.filename == '':
            print("DEBUG: Empty filename")
            return json.dumps({"error": "No file selected", "success": False})
            
        # Read file content and generate hash
        file_content = file.read()
        print(f"DEBUG: File content length: {len(file_content)}")
        file_hash = pdf_cache.get_file_hash(file_content)
        
        # Check if already cached
        if pdf_cache.is_cached(file_hash):
            cached_data = pdf_cache.get_chunks(file_hash)
            return json.dumps({
                "success": True,
                "file_hash": file_hash,
                "cached": True,
                "chunks_count": len(cached_data.get('chunks', [])),
                "processed_at": cached_data.get('processed_at', '').isoformat() if cached_data.get('processed_at') else ''
            })
        
        # Reset file pointer for processing
        file.seek(0)
        
        # Determine file type and process accordingly
        file_type = get_file_type(file.filename)
        
        start_time = time.time()
        
        print(f"Processing {file_type} file: {file.filename}")
        
        # Extract text based on file type
        text_content = extract_text_from_file(file, file_type)
        print(f"DEBUG: Extracted text content length: {len(text_content) if text_content else 'None'}")
        print(f"DEBUG: First 100 chars of text: {text_content[:100] if text_content else 'None'}")
        
        if text_content is None:
            return json.dumps({
                "error": f"Error extracting text from {file_type} file. Please ensure the file is valid and contains readable text.", 
                "success": False
            })
        
        if not text_content.strip():
            print("DEBUG: Text content is empty after stripping")
            return json.dumps({"error": "File appears to be empty", "success": False})
        
        print(f"File content length: {len(text_content)}")
        
        # Chunk the text first
        chunks = chunk_text(text_content)
        print(f"Created {len(chunks)} chunks")
        
        # Store chunks temporarily for validation
        pdf_cache.store_chunks(file_hash, text_content, chunks, chunks)
        
        # Return initial success with chunk info and trigger validation
        return json.dumps({
            "success": True,
            "file_hash": file_hash,
            "cached": False,
            "chunks_count": len(chunks),
            "text_length": len(text_content),
            "validation_required": True,
            "message": "Document chunked successfully. Validating document type..."
        })
        
        # Continue with processing - chunks already created above
        
        # Fast upload - minimal pre-processing for speed  
        simplified_chunks = []
        if OPENROUTER_API_KEY and len(chunks) <= 5:  # Only pre-simplify very small documents
            print(f"Pre-processing {len(chunks)} chunks (small document)...")
            for i, chunk in enumerate(chunks):
                try:
                    simplified_chunk = simplify_chunk_with_gemini(chunk)
                    simplified_chunks.append(simplified_chunk)
                    print(f"Pre-processed chunk {i+1}/{len(chunks)}")
                except Exception as e:
                    print(f"Error simplifying chunk {i+1}: {e}")
                    simplified_chunks.append(chunk)
        else:
            print(f"Fast upload mode - using original chunks ({len(chunks)} chunks)")
            simplified_chunks = chunks
        
        # Store in cache
        pdf_cache.store_chunks(file_hash, text_content, chunks, simplified_chunks)
        
        processing_time = time.time() - start_time
        
        return json.dumps({
            "success": True,
            "file_hash": file_hash,
            "cached": False,
            "chunks_count": len(chunks),
            "simplified_chunks_count": len(simplified_chunks),
            "processing_time": round(processing_time, 2),
            "text_length": len(text_content)
        })
        
    except Exception as e:
        print(f"Error in PDF upload: {e}")
        return json.dumps({"error": str(e), "success": False})

@app.route('/validate-document', methods=["POST"])
def validate_document():
    """Validate if the uploaded document is a legal document using AI"""
    print("Document validation endpoint called")
    
    try:
        data = request.get_json()
        if not data:
            return json.dumps({"error": "No JSON data provided", "success": False})
        
        file_hash = data.get('file_hash', '')
        
        if not file_hash:
            return json.dumps({"error": "No file hash provided", "success": False})
        
        # Check if file is cached
        if not pdf_cache.is_cached(file_hash):
            return json.dumps({"error": "File not found in cache. Please upload the file again.", "success": False})
        
        # Get cached data
        cached_data = pdf_cache.get_chunks(file_hash)
        chunks = cached_data.get('chunks', [])
        
        if not chunks:
            return json.dumps({"error": "No chunks found for validation", "success": False})
        
        # Validate document type using AI
        print(" Validating document type with AI...")
        is_legal, confidence, doc_type, explanation = validate_legal_document_with_ai(chunks)
        
        # Strict validation - reject if not legal or low confidence
        if not is_legal or confidence < 0.6:
            return json.dumps({
                "success": False,
                "is_legal_document": False,
                "confidence_score": round(confidence * 100, 1),
                "document_type": doc_type,
                "ai_explanation": explanation,
                "error": " Warning: This document does not appear to be a legal document.",
                "suggestion": "Please upload a legal document such as contracts, agreements, legal briefs, statutes, regulations, or other legal texts for proper analysis."
            })
        
        # Document is validated as legal
        print(f"Legal document validated - Type: {doc_type}")
        
        return json.dumps({
            "success": True,
            "is_legal_document": True,
            "confidence_score": round(confidence * 100, 1),
            "document_type": doc_type,
            "ai_explanation": explanation,
            "message": f" Document validated as: {doc_type}"
        })
        
    except Exception as e:
        print(f"Error in document validation: {e}")
        return json.dumps({"error": str(e), "success": False})

@app.route('/ask-question', methods=["POST"])
def ask_question_cached():
    """Ask question using cached PDF chunks"""
    print("Ask question endpoint called")
    
    try:
        data = request.get_json()
        if not data:
            return json.dumps({"error": "No JSON data provided", "success": False})
        
        file_hash = data.get('file_hash', '')
        question = data.get('question', '')
        
        if not file_hash:
            return json.dumps({"error": "No file hash provided", "success": False})
        
        if not question.strip():
            return json.dumps({"error": "No question provided", "success": False})
        
        # Check if file is cached
        if not pdf_cache.is_cached(file_hash):
            return json.dumps({"error": "File not found in cache. Please upload the file again.", "success": False})
        
        # Get cached data
        cached_data = pdf_cache.get_chunks(file_hash)
        text_chunks = cached_data.get('simplified_chunks', cached_data.get('chunks', []))
        
        if not text_chunks:
            return json.dumps({"error": "No chunks found for this file", "success": False})
        
        print(f"Processing question with {len(text_chunks)} cached chunks")
        
        # Use OpenRouter API
        if OPENROUTER_API_KEY:
            print("Using OpenRouter API for cached question...")
            # Check if we have simplified chunks
            has_simplified_chunks = 'simplified_chunks' in cached_data and cached_data['simplified_chunks']
            print(f"Using pre-simplified chunks: {has_simplified_chunks}")
            gemini_answer = query_gemini(question, text_chunks, chunks_already_simplified=has_simplified_chunks)
            answer = [{
                "answer": gemini_answer
            }]
            
            return json.dumps({
                "success": True,
                "answers": answer,
                "cached": True,
                "chunks_used": len(text_chunks)
            })
        else:
            return json.dumps({"error": "Gemini API not configured", "success": False})
            
    except Exception as e:
        print(f"Error in ask question: {e}")
        return json.dumps({"error": str(e), "success": False})




@app.route('/contracts', methods=["POST"])
def getContractResponse():
    print("Contract endpoint called")  # Debug log
    
    try:
        # Check if file is present in request
        if 'file' not in request.files:
            print("No file in request")
            return json.dumps([{"answer": "No file uploaded"}])
        
        file = request.files["file"]
        question = request.form.get('question', '')
        
        print(f"File received: {file.filename if file else 'None'}")
        print(f"File size: {len(file.read()) if file else 0} bytes")
        file.seek(0)  # Reset file pointer after reading for size
        print(f"Question received: {question}")
        
        # Check if file is empty
        if file.filename == '':
            print("Empty filename")
            return json.dumps([{"answer": "No file selected"}])
            
    except Exception as e:
        print(f"Error getting form data: {e}")
        return json.dumps([{"answer": f"Error processing form data: {str(e)}"}])

    # Process the file (text or PDF)
    try:
        file_extension = file.filename.lower().split('.')[-1] if file.filename else ''
        
        if file_extension == 'pdf':
            print("Processing PDF file...")
            # Reset file pointer
            file.seek(0)
            paragraph = extract_text_from_pdf(file)
            if paragraph is None:
                return json.dumps([{"answer": "Error extracting text from PDF"}])
        else:
            print("Processing text file...")
            # Read file content
            file_content = file.read()
            
            # Try different encodings
            try:
                paragraph = file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    paragraph = file_content.decode("latin-1")
                except UnicodeDecodeError:
                    paragraph = file_content.decode("cp1252")
        
        print(f"File content length: {len(paragraph)}")
        print(f"First 100 characters: {paragraph[:100]}")
        
        if not paragraph.strip():
            print("File appears to be empty")
            return json.dumps([{"answer": "File appears to be empty"}])
            
    except Exception as e:
        print(f"Error reading file: {e}")
        return json.dumps([{"answer": f"Error reading file: {str(e)}"}])

    # Validate inputs
    if not paragraph.strip():
        print("Empty paragraph content")
        return json.dumps([{"answer": "Document content is empty"}])
        
    if not question.strip():
        print("Empty question")
        return json.dumps([{"answer": "No question provided"}])

    if len(paragraph.strip()) > 0 and len(question.strip()) > 0:
        print('Processing document with AI...')
        print(f"Document length: {len(paragraph)} characters")
        print(f"Question: {question}")
        
        # Check if this is a request with cached data
        file_hash = request.form.get('file_hash', '')
        
        if file_hash and pdf_cache.is_cached(file_hash):
            print("Using cached chunks for processing")
            cached_data = pdf_cache.get_chunks(file_hash)
            text_chunks = cached_data.get('simplified_chunks', cached_data.get('chunks', []))
            print(f"Using {len(text_chunks)} cached chunks")
        else:
            # Chunk the text for better processing (legacy mode)
            text_chunks = chunk_text(paragraph)
            print(f"Document split into {len(text_chunks)} chunks")
        
        # Use Gemini API for processing
        print("Using Gemini API...")
        # Check if we're using cached simplified chunks
        chunks_simplified = file_hash and pdf_cache.is_cached(file_hash)
        gemini_answer = query_gemini(question, text_chunks, chunks_simplified)
        answer = [{
            "answer": gemini_answer
        }]
        
        return json.dumps(answer)

    else:
        return "Unable to call model, please select question and contract"

 




@app.route('/contracts/paraphrase/<path:selected_response>', methods=['GET'])
def getContractParaphrase(selected_response):
    print(selected_response)
    
    if selected_response == "":
        return "No answer found in document"
    else:
        print('getting paraphrases')
        paraphrases = paraphrase(selected_response)
        print(paraphrases)
        return paraphrases


@app.route('/get_response', methods=['POST'])
def get_response():
    question = request.form['selected_response']
    with open('responses.json', 'r') as file:
        responses = json.load('responses.json')
        for response in responses:
            if response['question'] == question:
                return response['answer']
    
    return "Response not found"

@app.route('/upload-multiple', methods=["POST"])
def upload_and_process_multiple():
    """Upload and process multiple documents, return combined hash for subsequent questions"""
    print("Multiple documents upload endpoint called")
    
    try:
        # Check if files are present in request
        if 'files' not in request.files:
            return json.dumps({"error": "No files uploaded", "success": False})
        
        files = request.files.getlist('files')
        
        if not files or len(files) == 0:
            return json.dumps({"error": "No files selected", "success": False})
            
        print(f"Processing {len(files)} files...")
        
        start_time = time.time()
        combined_text = ""
        processed_files = []
        total_chunks = []
        
        # Process each file
        for file in files:
            if file.filename == '':
                continue
                
            print(f"Processing file: {file.filename}")
            
            # Determine file type and extract text
            file_type = get_file_type(file.filename)
            text_content = extract_text_from_file(file, file_type)
            
            if text_content is None:
                print(f"Warning: Could not extract text from {file.filename}")
                continue
                
            if not text_content.strip():
                print(f"Warning: {file.filename} appears to be empty")
                continue
                
            # Chunk the text first for better analysis
            file_chunks = chunk_text(text_content)
            
            # Validate if the document is a legal document using AI
            print(f" Validating {file.filename} with AI...")
            is_legal, confidence, doc_type, explanation = validate_legal_document_with_ai(file_chunks)
            
            # Reject if AI determines it's not a legal document with high confidence
            if not is_legal and confidence >= 0.7:
                print(f" Warning: {file.filename} does not appear to be a legal document")
                return json.dumps({
                    "error": f" Warning: The file '{file.filename}' does not appear to be a legal document.",
                    "success": False,
                    "is_legal_document": False,
                    "confidence_score": round(confidence * 100, 1),
                    "document_type": doc_type,
                    "ai_explanation": explanation,
                    "problematic_file": file.filename,
                    "suggestion": "Please ensure all uploaded files are legal documents such as contracts, agreements, legal briefs, statutes, regulations, or other legal texts for proper analysis."
                })
            
            if not is_legal and confidence >= 0.5:  # Medium confidence warning but still proceed
                print(f" Possible non-legal document: {file.filename} - Type: {doc_type}")
            elif is_legal:
                print(f" {file.filename} validated as legal document - Type: {doc_type}")
            else:
                print(f" Uncertain document type: {file.filename} - Type: {doc_type}, proceeding anyway")
                
            # Add file separator and combine text
            combined_text += f"\n\n=== DOCUMENT: {file.filename} ===\n\n"
            combined_text += text_content
            processed_files.append(file.filename)
            
            # Reset file pointer for next operations
            file.seek(0)
        
        if not combined_text.strip():
            return json.dumps({"error": "No readable content found in uploaded files", "success": False})
        
        # Generate hash for combined content
        combined_hash = pdf_cache.get_file_hash(combined_text.encode('utf-8'))
        
        # Check if already cached
        if pdf_cache.is_cached(combined_hash):
            cached_data = pdf_cache.get_chunks(combined_hash)
            return json.dumps({
                "success": True,
                "combined_hash": combined_hash,
                "cached": True,
                "total_chunks": len(cached_data.get('chunks', [])),
                "files_processed": len(processed_files),
                "processed_at": cached_data.get('processed_at', '').isoformat() if cached_data.get('processed_at') else ''
            })
        
        print(f"Combined content length: {len(combined_text)}")
        
        # Chunk the combined text
        chunks = chunk_text(combined_text)
        print(f"Created {len(chunks)} chunks from {len(processed_files)} files")
        
        # Fast upload - minimal pre-processing for speed  
        simplified_chunks = []
        if OPENROUTER_API_KEY and len(chunks) <= 10:  # Only pre-simplify small documents
            print(f"Pre-processing {len(chunks)} chunks (small combined document)...")
            for i, chunk in enumerate(chunks):
                try:
                    simplified_chunk = simplify_chunk_with_gemini(chunk)
                    simplified_chunks.append(simplified_chunk)
                    print(f"Pre-processed chunk {i+1}/{len(chunks)}")
                except Exception as e:
                    print(f"Error simplifying chunk {i+1}: {e}")
                    simplified_chunks.append(chunk)
        else:
            print(f"Fast upload mode - using original chunks ({len(chunks)} chunks)")
            simplified_chunks = chunks
        
        # Store in cache
        pdf_cache.store_chunks(combined_hash, combined_text, chunks, simplified_chunks)
        
        processing_time = time.time() - start_time
        
        return json.dumps({
            "success": True,
            "combined_hash": combined_hash,
            "cached": False,
            "total_chunks": len(chunks),
            "files_processed": len(processed_files),
            "processing_time": round(processing_time, 2),
            "total_text_length": len(combined_text),
            "processed_files": processed_files
        })
        
    except Exception as e:
        print(f"Error in multiple files upload: {e}")
        return json.dumps({"error": str(e), "success": False})


if __name__ == '__main__':
    # Get port from environment variable or use default
    port = int(os.environ.get('FLASK_RUN_PORT', 5000))
    print(f"Starting Flask server on port {port}")
    app.run(host='127.0.0.1', port=port, debug=True)